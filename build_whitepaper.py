from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/chenglitao/Desktop/work_project/二元期权量化/事件合约系统化交易白皮书_讨论稿_v0.1.docx")

FONT_CN = "Maple Mono NF CN"
FONT_LATIN = "Maple Mono NF CN"
NAVY = RGBColor(32, 55, 72)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GOLD = RGBColor(122, 90, 0)
RED = RGBColor(155, 28, 28)
GREEN = RGBColor(35, 92, 67)
GRAY = RGBColor(86, 94, 103)
LIGHT_GRAY = "F4F6F9"
LIGHT_BLUE = "E8EEF5"
LIGHT_GOLD = "FFF8E8"
LIGHT_RED = "FDECEC"
WHITE = RGBColor(255, 255, 255)


def set_run_font(run, size=11, bold=None, italic=None, color=None, name_cn=FONT_CN, name_latin=FONT_LATIN):
    run.font.name = name_latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name_latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name_latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name_cn)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def apply_table_geometry(table, widths_dxa: Sequence[int], indent_dxa=120):
    total = sum(widths_dxa)
    if total != 9360:
        raise ValueError(f"table widths must sum to 9360 DXA, got {total}")
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="D6DAE1", size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def make_numbering(doc: Document):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abs or [0]) + 1
    next_num = max(existing_num or [0]) + 1

    def add_abstract(abstract_id, fmt, text, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "279")
        p_pr.append(ind)
        lvl.append(p_pr)
        if font:
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), font)
            r_fonts.set(qn("w:hAnsi"), font)
            r_fonts.set(qn("w:eastAsia"), font)
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    add_abstract(next_abs, "bullet", "•", FONT_LATIN)
    bullet_num = next_num
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(bullet_num))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(next_abs))
    num.append(abs_id)
    numbering.append(num)

    add_abstract(next_abs + 1, "decimal", "%1.")
    decimal_num = next_num + 1
    num2 = OxmlElement("w:num")
    num2.set(qn("w:numId"), str(decimal_num))
    abs_id2 = OxmlElement("w:abstractNumId")
    abs_id2.set(qn("w:val"), str(next_abs + 1))
    num2.append(abs_id2)
    numbering.append(num2)
    return bullet_num, decimal_num


def set_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def add_para(doc, text="", *, size=11, bold=False, italic=False, color=None,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=8, line=1.333,
             keep_with_next=False):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep_with_next
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_rich_para(doc, runs, **kwargs):
    p = doc.add_paragraph()
    p.alignment = kwargs.pop("align", WD_ALIGN_PARAGRAPH.JUSTIFY)
    pf = p.paragraph_format
    pf.space_before = Pt(kwargs.pop("before", 0))
    pf.space_after = Pt(kwargs.pop("after", 8))
    pf.line_spacing = kwargs.pop("line", 1.333)
    for item in runs:
        if isinstance(item, str):
            text, opts = item, {}
        else:
            text, opts = item
        r = p.add_run(text)
        set_run_font(r, size=opts.get("size", 11), bold=opts.get("bold"),
                     italic=opts.get("italic"), color=opts.get("color"))
    return p


def add_bullet(doc, text, num_id, *, bold_prefix=None, color=None):
    p = doc.add_paragraph()
    set_num(p, num_id)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.208
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=color)
    return p


def add_number(doc, text, num_id):
    p = doc.add_paragraph()
    set_num(p, num_id)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.208
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_callout(doc, label, text, *, fill=LIGHT_BLUE, color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    apply_table_geometry(table, [9360])
    set_table_borders(table, color=fill, size=8)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    r1 = p.add_run(label + "  ")
    set_run_font(r1, size=11, bold=True, color=color)
    r2 = p.add_run(text)
    set_run_font(r2, size=11, color=NAVY)
    add_para(doc, "", after=4, line=1.0)
    return table


def add_table(doc, headers, rows, widths_dxa, *, header_fill=LIGHT_BLUE,
              font_size=9.4, center_cols: Iterable[int] = ()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    apply_table_geometry(table, widths_dxa)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    set_row_cant_split(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.12
        r = p.add_run(str(header))
        set_run_font(r, size=font_size, bold=True, color=NAVY)
    for row in rows:
        data_row = table.add_row()
        set_row_cant_split(data_row)
        cells = data_row.cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in center_cols else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=RGBColor(35, 39, 44))
    apply_table_geometry(table, widths_dxa)
    add_para(doc, "", after=4, line=1.0)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=9, color=GRAY)


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    h_tokens = {
        1: (16, BLUE, 18, 10),
        2: (13, BLUE, 12, 6),
        3: (12, DARK_BLUE, 8, 4),
    }
    for level, (size, color, before, after) in h_tokens.items():
        st = styles[f"Heading {level}"]
        st.font.name = FONT_LATIN
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("事件合约系统化交易｜自营研究白皮书（讨论稿）")
    set_run_font(hr, size=9, color=GRAY)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def cover(doc):
    add_para(doc, "", after=80, line=1.0)
    add_para(doc, "独立研究项目白皮书", size=11, bold=True, color=GOLD,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=18, line=1.0)
    add_para(doc, "事件合约系统化交易", size=30, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=8, line=1.0)
    add_para(doc, "Systematic Event-Contract Trading", size=15, color=DARK_BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=3, line=1.0)
    add_para(doc, "夫妻小团队的清洁室复刻、研究与小资金验证方案", size=13, color=DARK_BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=30, line=1.15)
    add_para(doc, "讨论稿 v0.1｜2026年8月17日", size=11, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=8, line=1.0)
    add_para(doc, "面向：金融研究背景 × 计算机工程背景的家庭自营研究团队", size=9.5,
             italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=70, line=1.1)
    add_callout(doc, "使用边界", "本文件用于家庭内部认知对齐、工程规划与研究决策，不构成法律、投资或开户建议。任何实盘前均须完成居住地法律、平台资格、雇主合规及资金路径审查。", fill=LIGHT_GOLD, color=GOLD)
    add_para(doc, "核心建议：条件式 GO —— 可以先做公开数据、独立实现和模拟交易；不建议在合规审查、回测与影子交易完成前投入全部 3,000 美元。",
             size=10.5, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, before=16, after=0, line=1.2)
    doc.add_page_break()


def main():
    doc = Document()
    configure_doc(doc)
    bullet_num, decimal_num = make_numbering(doc)
    doc.core_properties.title = "事件合约系统化交易白皮书（讨论稿）"
    doc.core_properties.subject = "预测市场与事件合约自营量化研究项目"
    doc.core_properties.author = "家庭自营研究团队"
    doc.core_properties.keywords = "预测市场, 事件合约, 量化交易, 回测, 低延迟, 风险管理"
    cover(doc)

    add_heading(doc, "文档定位与决策问题", 1)
    add_para(doc, "本白皮书讨论的不是是否复制他人的代码或策略，而是能否仅基于公开资料、公开 API 和独立工程，实现一套可验证的事件合约系统化交易能力。项目目标应定义为“能力复刻”，而不是“收益曲线复刻”。前者可通过数据、仿真、执行和风控验收；后者受策略细节、延迟位置、资金路径和运气影响，无法从演示系统直接推导。")
    add_callout(doc, "建议决策", "批准 6 个月研究立项；阶段 0–2 限于数据、回测和影子交易。只有法律/雇主合规、账户资格、样本外收益和执行偏差四类门槛同时通过，才进入最高 300 美元的首轮实盘。", fill=LIGHT_BLUE, color=DARK_BLUE)
    add_heading(doc, "执行摘要", 1)
    for item in [
        "赛道名称：建议采用“事件合约系统化交易”（Systematic Event-Contract Trading，SECT）；对外可用“预测市场量化交易”。避免以“二元期权量化”作为主名称。",
        "赛道判断：高速增长、数据开放、市场碎片化、规则复杂、容量有限。2025 年 CFTC 监管预测市场交易量已超过 250 亿美元；第三方数据估计 2026 年 7 月 Kalshi、Polymarket 与 Polymarket US 合计交易量约 506 亿美元，但世界杯等大型事件造成明显周期波动。[S05][S24]",
        "Demo 判断：SignalX 已完成实时采集、ClickHouse 数仓、策略研究、低延迟引擎、账户/KMS、版本化配置、日志监控和实盘 PnL 闭环；其工程真实性较强，但市场数据前端、数据质量、PnL 锁账和回测真实性仍处于早期阶段。[S01][S02]",
        "Alpha 判断：短周期加密事件的现货—事件合约 lead-lag、结构性逻辑套利、尾盘/结算策略、跨平台套利、做市与激励、行为偏差和公开交易者研究均有空间。夫妻店应优先做可离线验证、对极致延迟依赖较低的路线。",
        "工程判断：真正难点不是下单 API，而是市场主数据、订单簿重建、事件驱动回测、延迟/排队/部分成交仿真、不可修改账本、风控与调账。MVP 约需 500–900 人小时；稳定运行通常需要 6–12 个月。",
        "资金判断：3,000 美元适合作为最终风险资本上限，不适合第一天全部入金。建议按 0 → 100 → 300 → 1,000 → 3,000 美元递进；研究基础设施预算应单列。",
        "最大阻断项：中国大陆的虚拟货币、境外平台和网络竞猜法律风险，以及配偶所在量化公司的个人交易/副业/IP 规定。没有书面或专业确认前，只做研究、数据和模拟。",
    ]:
        add_bullet(doc, item, bullet_num)

    add_heading(doc, "阅读路线", 2)
    add_table(doc,
              ["章节", "回答的问题", "建议讨论人"],
              [
                  ("0", "这条赛道叫什么，与股票量化差别在哪里？", "双方"),
                  ("1", "哪些是已观察事实，哪些只是宣传或待验证？", "双方"),
                  ("2", "可研究的 Alpha 有哪些，优先级如何？", "量化研究侧牵头"),
                  ("3", "系统全貌、回测要求、设备与云成本是什么？", "工程侧牵头"),
                  ("4", "夫妻店如何分工、用多少钱、如何分阶段？", "双方"),
                  ("5", "哪些风险会让项目立即停止？", "双方＋外部合规"),
              ], [1100, 5060, 3200], center_cols=(0,))

    add_heading(doc, "0. 赛道定义、特征与命名", 1)
    add_heading(doc, "0.1 推荐名称", 2)
    add_para(doc, "推荐正式名称为“事件合约系统化交易”（SECT）。它强调交易对象是由未来事件结果结算的合约，方法是系统化研究、执行与风险控制。若突出行业，可称“预测市场量化交易”；若突出当前研究重点，可称“预测市场微观结构与低延迟交易”。")
    add_table(doc, ["名称", "适用场景", "评价"], [
        ("事件合约系统化交易", "白皮书、内部项目、正式研究", "最准确；兼容受监管与链上平台"),
        ("预测市场量化交易", "对外沟通、行业研究", "通俗，但容易被理解为只做观点预测"),
        ("事件驱动高频交易", "强调低延迟时", "容易和股票新闻事件交易混淆"),
        ("二元期权量化", "不建议作为主名称", "在中国语境下可能指庄家式押涨跌平台，法律与声誉风险高"),
    ], [2300, 3100, 3960])
    add_heading(doc, "0.2 与股票量化赛道的本质区别", 2)
    add_table(doc, ["维度", "事件合约/预测市场", "股票/期货量化"], [
        ("收益结构", "0/1 或有限离散结果；可能本金归零", "连续价格变化，通常可动态退出"),
        ("标的寿命", "大量一次性、短寿命市场", "证券主数据相对稳定、长期存在"),
        ("规则", "标题不足以定义资产；结算来源、截止时间、边界条款属于定价核心", "交易所规则重要，但单个证券通常不因自然语言规则而改变结算"),
        ("风险分布", "高胜率与罕见全损并存；尾部风险显著", "可呈多种分布，但不天然二元"),
        ("数据", "公开 API、链上成交和地址活动较丰富；撤单/报价生命周期可能缺失", "专业逐笔数据昂贵，但结构和时间戳规范较成熟"),
        ("市场结构", "平台碎片化、订单簿较薄、结算机制各异", "主流交易所深度更高，规则与基础设施更标准"),
        ("Alpha 半衰期", "事件与平台特定，迁移性弱，容量小", "部分因子跨标的可复用，容量相对更大"),
        ("合规", "金融、博彩、虚拟货币、地域限制交叉", "证券/期货监管路径相对清晰"),
    ], [1700, 3830, 3830], font_size=8.8)
    add_callout(doc, "赛道共性", "定价不是单纯预测事件，而是同时预测：真实世界结果、平台采用的结算定义、订单簿如何反应，以及自己的订单能否在价格失效前成交。", fill=LIGHT_GOLD, color=GOLD)

    add_heading(doc, "1. 已知事实、背景事实与已验证点", 1)
    add_heading(doc, "1.1 SignalX Demo 的可观察事实", 2)
    add_para(doc, "以下数据来自 2026 年 8 月 17 日对 SignalX 控制台和项目介绍的只读检查。它们可证明系统在运行，但不能代替平台对账单、链上地址、资金流水或第三方审计。[S01][S02]")
    add_table(doc, ["维度", "观察值", "可得结论", "限制"], [
        ("数仓", "约 148.1 亿行；386GB 压缩；磁盘使用 2.5TB；43 张表", "已存在持续采集与查询体系", "统计口径和历史完整性未独立复核"),
        ("实时性", "最新分片更新到当晚约 19:16", "采集在持续运行", "覆盖率监控查询因旧表名报错"),
        ("Agent", "17 个登记进程，5 个在线", "存在多地域/多用途守护进程", "其余进程是停用、备用还是故障未知"),
        ("配置", "10 个启用配置；TOML 不可变版本", "配置治理思路较专业", "版本审批和回滚 SLA 未知"),
        ("PnL", "8 月 +$915.17；1,270 条已结算；12 盈利日/5 亏损日", "短样本内已产生实盘盈利", "未核对手续费、返佣、存取款、未结算仓位"),
        ("资金集中", "Predict.fun 主账户贡献约 93.5% 月度盈利", "当前优势高度集中", "不能据此证明跨平台能力"),
        ("日志", "实时出现订阅、信号、下单、撤单、结算方向和订单 ID", "不是静态 PPT 系统", "存在 stale market、重连与状态机警告"),
    ], [1300, 2500, 2700, 2860], font_size=8.2)
    add_heading(doc, "1.2 已跑通的验证点", 2)
    for item in [
        "公开和平台行情可以被稳定采集、压缩和关联查询。",
        "短周期加密事件合约可以与 Binance/Chainlink 等参考价格连接。",
        "策略信号可以经过账户授权、预交易约束、下单、成交追踪和结算进入统一 PnL。",
        "在当前平台和短样本期内，Winner/Tail Sweep 类策略至少产生过正的已结算 PnL。",
        "多机房部署、连接预热和低延迟工程对结算附近的竞争具有实际价值。",
        "版本化配置、KMS、日志和控制台是长期运行所必需，而不是额外包装。",
    ]:
        add_bullet(doc, item, bullet_num)
    add_heading(doc, "1.3 尚未验证或存在反证的点", 2)
    for item in [
        "长期稳健性：观察期只有约三周；8 月 1 日单日盈利占月度盈利约 53%。",
        "资本可扩展性：盘口深度、对手行为和平台限额可能使收益无法线性放大。",
        "回测真实性：项目材料自己把“拟合真实订单簿与订单流延迟”列为未来任务。",
        "PnL 可审计性：早期截图与当前控制台的多日历史数值发生修订，至 8 月 6 日累计差约 $34.41。",
        "8–9ms 优势：可以测量，但“快于 99.999% 对手”缺少公开样本、测量边界和对照组。",
        "跨平台能力：当前月度盈利全部来自 Predict.fun，其他账户为 0 或停用。",
    ]:
        add_bullet(doc, item, bullet_num, color=RED)

    add_heading(doc, "1.4 行业与平台背景", 2)
    add_para(doc, "CFTC 将预测市场产品称为 event contracts，通常以二元或多结果形式结算，用于投机、信息聚合或风险对冲。2026 年美国仍处于快速规则演进期，监管机构同时强调市场创新、市场操纵、内幕信息、合同规则和交易所自律义务。[S03][S04][S05]")
    add_para(doc, "主流平台形成两类架构：一类是 Kalshi 等受监管交易所；另一类是 Polymarket 等以中心化订单簿撮合、链上持仓与结算为特点的平台。Polymarket 的公共接口提供订单簿、成交、用户活动、持仓与解析事件；交易状态还经历 MATCHED、MINED、CONFIRMED、RETRYING 或 FAILED，因此“看到成交”不一定等于最终确认。[S06][S07][S10]")
    add_para(doc, "行业规模已足以支持专业参与者，但仍小于成熟期货市场。CFTC 文件称 2025 年受监管预测市场交易量超过 250 亿美元；第三方数据估计 2026 年 7 月三大平台合计约 506 亿美元，其中世界杯造成显著放量。规模增长不等于策略容量增长，尤其是单一短周期盘口。[S05][S24]")
    add_heading(doc, "1.5 公开数据的价值与盲区", 2)
    add_table(doc, ["可获得", "研究用途", "关键盲区"], [
        ("市场规则、截止时间、结果", "市场主数据、规则解析、结算研究", "规则可能补充或争议；标题不等于规则"),
        ("L2 订单簿和价格变更", "重建盘口、价差、深度、短期反应", "聚合价位不一定包含精确队列位置"),
        ("成交和链上地址活动", "交易者画像、资金流、持仓与已实现收益", "公开成交缺少私人信号和完整决策过程"),
        ("自己的订单/成交 WebSocket", "OMS、状态机、调账和执行质量", "需要安全保存凭证，且存在重试/最终性"),
        ("参考现货与预言机", "lead-lag、公允概率和结算源对齐", "平台采用的价格源、采样窗口和精度必须逐市场确认"),
    ], [2200, 3200, 3960], font_size=8.8)
    add_para(doc, "2026 年两项 Polymarket 微观结构研究给出重要警告：仅凭公开报价流推断买卖方向，和链上真实成交方向的一致度可能只有约 59%；同时，链下订单放置和撤单生命周期无法从链上地址完整恢复，因此从成交记录“复制别人策略”存在结构性不可识别问题。[S15][S16]")

    add_heading(doc, "2. Alpha 路线：现有路线与可拓展方向", 1)
    add_heading(doc, "2.1 Alpha 分类框架", 2)
    add_table(doc, ["路线", "经济来源", "工程依赖", "适合夫妻店", "建议阶段"], [
        ("现货—事件合约 Lead-Lag", "参考现货先动，事件盘口后反应", "高质量时间戳、实时行情、成本模型", "高", "第一优先"),
        ("逻辑/结构套利", "YES/NO 或互斥结果的总价格异常", "市场规则图谱、原子或容错执行", "高", "第一优先"),
        ("结算/Tail Sweep", "结果趋明但订单簿未完全收敛", "极低延迟、精确结算源、强风控", "中低", "第二阶段"),
        ("跨平台套利", "同一经济事件跨所定价偏差", "双账户、资本分仓、规则映射", "中", "第二阶段"),
        ("做市/激励", "点差、返佣与流动性奖励", "库存模型、撤单可靠性、队列仿真", "中", "第二阶段"),
        ("行为与校准偏差", "散户概率误判、长赔偏差、热门事件偏差", "大样本、严格样本外检验", "高", "研究并行"),
        ("公开交易者行为", "识别高水平交易者或条件策略", "地址聚类、交易重建、因果辨识", "中", "研究并行"),
        ("新闻/LLM 事件定价", "文本信息转概率快于市场", "信息源、评测集、低幻觉输出", "中低", "后期"),
    ], [1900, 2500, 2200, 1100, 1660], font_size=8.0, center_cols=(3, 4))

    add_heading(doc, "2.2 现货—事件合约 Lead-Lag", 2)
    add_para(doc, "这是最适合独立复刻的基准路线。对 BTC/ETH/BNB 的 5 分钟、15 分钟涨跌合约，使用多个现货源、指数或预言机计算在剩余期限内收于阈值上方的概率，并测量事件盘口相对现货的响应延迟。它既可以形成 taker 信号，也可以形成 maker 报价。")
    for item in [
        "输入：现货逐笔、BBO、事件盘口、波动率、剩余时间、平台结算源。",
        "最小模型：数字期权概率 + 短时波动率 + 盘口深度；复杂模型可加入跳跃、微观价格和多源延迟。",
        "验收：在样本外、扣除费用和最坏合理滑点后，按事件日分层仍为正；收益不由少数结算窗口贡献。",
        "风险：对手使用不同模型不代表错误；现货价格先动也可能只是噪声。",
    ]:
        add_bullet(doc, item, bullet_num)

    add_heading(doc, "2.3 逻辑与结构套利", 2)
    add_para(doc, "二元合约天然存在概率约束，例如同一条件下 YES 与 NO 的可执行价格之和、互斥且完备结果集合的价格之和、嵌套事件的单调性。偏离并不自动等于无风险套利：手续费、深度、部分成交、持仓铸造/合并和结算规则差异都会破坏理论利润。")
    add_callout(doc, "优点", "信号解释清楚、容易离线验证、适合构建回测基线；即使最终不能盈利，也能检验市场主数据、规则解析、订单簿和执行系统。", fill=LIGHT_BLUE, color=DARK_BLUE)

    add_heading(doc, "2.4 结算与 Tail Sweep", 2)
    add_para(doc, "Demo 的主策略是在结果即将确定或已经可以从参考源推断时，以接近 1 的价格买入胜出方向。其表面胜率很高，但风险收益高度非对称：以 0.99 买入，正常一次只赚约 0.01；一次完整判断错误，理论上可抵消约 99 次正常利润，尚未计费用、重复下单和平台故障。")
    for item in [
        "必须独立维护平台的结算源、采样窗口、精度、时区和争议规则。",
        "信号要分为“方向推断”“官方可确认”“平台已解析”，不同阶段采用不同限额。",
        "任何数据停滞、时钟偏移、价格源分叉、规则不确定都必须 fail closed。",
        "3,000 美元账户不得照搬 Demo 中单市场 50–700 美元的名义风险。",
    ]:
        add_bullet(doc, item, bullet_num, color=RED)

    add_heading(doc, "2.5 跨平台套利", 2)
    add_para(doc, "跨平台价差来自参与者结构、规则、结算资产和网络位置差异。真正的工程问题是“同一经济事件”不一定是“同一法律与结算合约”。需要把每个市场映射为规范化事件、结果集合、结算来源和截止条件，并设置规则相似度阈值。")
    add_para(doc, "对 3,000 美元资本，跨所会把资金分散到多个账户，降低每边可用深度，同时增加 KYC、出入金和资金冻结风险。更适合作为数据研究与告警工具，待单平台策略和对账稳定后再实盘。")

    add_heading(doc, "2.6 做市与激励", 2)
    add_para(doc, "Polymarket 对多类市场收取 taker 费用并将部分费用以 maker rebate 返还；Kalshi 也存在流动性和做市激励。激励是 Alpha 的组成部分，但规则可随时变化，不能成为策略的唯一利润来源。[S08][S09][S12]")
    add_para(doc, "做市回测必须模拟排队位置、订单存活时间、部分成交、撤单在途和毒性流。管理科学关于限价单的研究强调，队列位置本身具有不确定性；把“盘口触价”当作自己成交，会系统性夸大收益。[S18]")

    add_heading(doc, "2.7 行为偏差与公开交易者研究", 2)
    add_para(doc, "长期存在的候选异常包括 favourite–longshot bias，即低概率结果被高估、高概率结果被低估，但其方向和强度随市场结构变化，且交易成本可能吞噬优势。[S14] 对公开交易者，更合理的目标是识别行为类型和条件信号，而不是逐笔复制。")
    for item in [
        "可做：按持仓周期、市场类别、成交频率、价格区间、结果校准和回撤构建交易者画像。",
        "不可直接做：把公开地址的下一笔成交视为可同步执行的信号；可见时往往已经过时。",
        "需防止：幸存者偏差、地址拆分、资金转移、复制者挤压和排行榜的已实现/未实现口径差异。",
        "LLM 用途：提出可编码假设、生成特征说明、审查实验日志；不负责决定是否上线。",
    ]:
        add_bullet(doc, item, bullet_num)

    add_heading(doc, "2.8 研究优先级", 2)
    add_table(doc, ["优先级", "首批实验", "成功标准", "失败也能沉淀"], [
        ("P0", "现货—事件盘口响应曲线", "样本外净 edge；时间戳与延迟可解释", "统一时间轴、行情与盘口重建"),
        ("P0", "YES/NO 与结果集合约束", "可执行而非中间价套利；部分成交可控", "市场主数据与规则图谱"),
        ("P1", "Tail Sweep 离线复盘", "极端错误率、结算分叉和断流压力测试通过", "结算源适配与故障注入"),
        ("P1", "做市/激励影子策略", "扣除逆向选择后为正；库存受控", "队列模型和 OMS"),
        ("P2", "公开地址行为聚类", "稳定的跨期行为类型，而非仅拟合收益", "研究数据集和实体解析"),
        ("P3", "LLM 新闻与自动因子", "严格冻结评测集；优于简单基线", "Agentic 研究平台"),
    ], [1000, 2800, 3300, 2260], font_size=8.5, center_cols=(0,))

    add_heading(doc, "3. 工程难度与 IT 投入估算", 1)
    add_heading(doc, "3.1 推测的完整系统全貌", 2)
    add_para(doc, "根据 Demo 页面、Agent 日志、数据表与配置治理，可推测其系统并非一个交易脚本，而是由研究面、交易面和共享控制面组成。独立复刻不需要一开始达到 Demo 的 148 亿行与多机房规模，但系统边界应从第一天设计正确。")
    architecture_rows = [
        ("数据接入", "预测市场 L2/成交/市场状态；Binance/指数/预言机；时间同步与断线恢复"),
        ("不可修改原始层", "原始 WebSocket/REST 消息、接收时间、源时间、序列号，按对象存储分区"),
        ("标准化数仓", "市场主数据、订单簿快照/增量、成交、参考价格、规则与解析结果；ClickHouse/Parquet"),
        ("研究与特征", "Notebook/批任务、实验注册、特征版本、标签、样本切分与结果归档"),
        ("事件驱动回测", "历史回放、订单簿重建、网络延迟、队列、部分成交、费用、结算与故障注入"),
        ("实时策略引擎", "同一策略接口支持 backtest/paper/live；只读信号与交易执行解耦"),
        ("OMS/EMS", "订单状态机、幂等、重试、撤单、成交最终性、持仓和账户适配"),
        ("风险引擎", "下单前限额、市场/策略/账户敞口、每日止损、断流 kill switch、人工总开关"),
        ("账本与调账", "成交、费用、返佣、铸造/合并、赎回、存取款和 NAV；每日不可修改快照"),
        ("控制与运维", "KMS、配置版本、部署、指标、日志、告警、故障演练、审计记录"),
    ]
    add_table(doc, ["系统层", "最低要求"], architecture_rows, [2100, 7260], font_size=8.8)

    add_heading(doc, "3.2 回测系统的专业门槛", 2)
    add_para(doc, "股票日频回测框架不能直接移植。事件合约的策略收益往往集中在秒级或结算边界，任何“按中间价立即成交”的假设都会严重失真。金融回测研究也表明，大量尝试后选择最好结果容易产生统计幻觉，应记录全部试验并评估回测过拟合概率。[S17]")
    add_table(doc, ["模块", "必须模拟", "常见错误"], [
        ("市场时间", "源时间、接收时间、处理时间、下单/确认时间", "使用单一时间戳；未来数据泄漏"),
        ("订单簿", "快照＋增量、序列缺口、跨价、最小 tick", "只用 OHLC 或 BBO"),
        ("成交", "队列位置、部分成交、FOK/FAK、撤单在途", "触价即全额成交"),
        ("费用", "taker fee、maker rebate、Gas/桥接、入金成本", "使用静态统一费率"),
        ("结算", "规则版本、参考源、争议、延迟、0/0.5/1 等结果", "只按标题或现货终值判断"),
        ("样本", "按事件/日期 purging、walk-forward、冻结样本外", "随机 K-fold；参数反复看测试集"),
        ("风险", "并发市场、重复订单、源分叉、平台宕机", "单笔独立、无故障的理想仿真"),
    ], [1500, 3900, 3960], font_size=8.5)
    add_callout(doc, "最低回测标准", "同一份策略代码或同一策略 DSL 同时运行于历史回放、实时影子和小资金实盘；差异只能来自明确的环境适配层。", fill=LIGHT_BLUE, color=DARK_BLUE)

    add_heading(doc, "3.3 与量化公司常见系统的对齐清单", 2)
    for item in [
        "Security/Market Master：事件、结果、平台合约、结算规则和版本化映射。",
        "Market Data Plant：多源接入、序列校验、时间同步、原始消息留存和质量监控。",
        "Research Platform：数据集版本、特征库、实验注册、参数搜索和样本外冻结。",
        "Exchange Simulator：订单簿、延迟、队列、费用、部分成交与故障模型。",
        "OMS/EMS：订单生命周期、幂等、重试、撤单、成交最终性和账户适配。",
        "Risk/PMS：预交易限额、持仓、相关敞口、日内风险、压力测试和 kill switch。",
        "Ledger/Reconciliation：平台账单、链上交易、内部账本和 NAV 对账。",
        "SRE/Security：KMS、密钥轮换、部署审批、日志指标、值班和灾备。",
    ]:
        add_bullet(doc, item, bullet_num)
    add_para(doc, "与对象沟通时，应讨论这些系统的通用验收原则，而不是其雇主的具体实现、代码、参数、架构图或策略。")

    add_heading(doc, "3.4 工程难度与人时", 2)
    add_table(doc, ["工作包", "难度(1–5)", "MVP 人时", "稳定化人时", "主要失败模式"], [
        ("市场主数据与规则", "4", "50–90", "80–160", "映射错误造成错误结算"),
        ("行情采集与数仓", "4", "80–140", "100–220", "增量缺口、重连重复、时钟漂移"),
        ("订单簿回放/回测", "5", "120–220", "200–400", "虚假成交、泄漏、成本漏计"),
        ("策略与研究工具", "3", "80–160", "持续投入", "实验不可复现、过拟合"),
        ("执行/OMS", "5", "100–180", "180–350", "重复订单、状态机分叉、撤单失败"),
        ("风险/账本/对账", "5", "80–150", "160–300", "PnL 漂移、敞口失真、无法审计"),
        ("部署/监控/密钥", "4", "60–100", "100–200", "凭证泄露、无告警、恢复失败"),
    ], [1900, 1200, 1300, 1500, 3460], font_size=8.3, center_cols=(1, 2, 3))
    add_para(doc, "去重后，夫妻店 MVP 预计 500–900 人小时；稳定化和跨平台通常超过 1,000 人小时。若两人合计每周投入 20–30 小时，形成可信 MVP 约需 4–6 个月；达到可长期无人值守通常需 6–12 个月。")

    add_heading(doc, "3.5 IT 设备与云成本", 2)
    add_para(doc, "成本按 2026 年公开价格和 Demo 量级推算。R2 标准存储约 $0.015/GB/月且公网出口免费；AWS 计算按需按秒计费，东京 c7i.large 的公开价格约 $0.11–0.12/小时量级。实际网络出口、磁盘 IOPS 和多区域流量可能高于计算实例本身。[S19][S20]")
    add_table(doc, ["阶段", "建议配置", "一次性投入", "月度云成本", "说明"], [
        ("本地研究", "16+ 核 CPU；64–128GB RAM；4–8TB NVMe；UPS", "已有设备：0；补强约 ¥3k–20k", "$10–60", "ClickHouse、Parquet、回测与备份"),
        ("单区域采集", "东京或伦敦 2–4 vCPU；100–300GB SSD", "0", "$100–250", "研究与低频影子交易"),
        ("单区域小实盘", "采集/交易分进程；监控与备用机", "0", "$200–500", "不含法律、入金和人工"),
        ("两区域专业化", "东京＋伦敦/美国；集中数仓；对象存储", "0", "$500–1,500", "接近小型专业系统，不必首期建设"),
        ("Demo 量级", "4 地域、多 Agent、TB 级磁盘、控制台/KMS", "本地服务器另计", "$800–2,500+", "仅为推测；没有必要一开始复制规模"),
    ], [1500, 2650, 1700, 1400, 2110], font_size=8.1, center_cols=(2, 3))
    add_callout(doc, "预算原则", "把 3,000 美元定义为“风险资本”，把设备、云、法律咨询和时间成本单独列账。若 3,000 美元是项目全部现金预算，则前 3–4 个月应只做数据与模拟，不应同时承担专业云成本和实盘风险。", fill=LIGHT_GOLD, color=GOLD)

    add_heading(doc, "4. 资金、人力与中短期项目规划", 1)
    add_heading(doc, "4.1 夫妻店角色设计", 2)
    add_table(doc, ["角色", "主要责任", "不得越界"], [
        ("工程负责人（你）", "数据接入、数仓、回测框架、OMS、风险、部署与审计；项目经理", "不根据收益曲线随意改策略；不接触对象雇主资产"),
        ("量化研究负责人（对象）", "假设、统计检验、成本模型、组合与风险评审；研究规范", "不带出雇主代码、数据、参数、未公开流程或客户信息"),
        ("共同职责", "合规门槛、上线审批、资金拨付、每日复核、停机决策", "任何一方均可一票停止实盘"),
    ], [2000, 4500, 2860], font_size=8.8)
    add_heading(doc, "4.2 六个月阶段计划与闸门", 2)
    add_table(doc, ["阶段", "时间", "交付物", "通过条件", "资金上限"], [
        ("−1 合规与清洁室", "第 1–2 周", "法律问题清单、雇主申报、平台资格、IP 边界", "无账户代持、无 VPN 绕过、无雇主冲突；律师/合规意见允许继续", "$0"),
        ("0 数据底座", "第 3–6 周", "单平台＋单参考源；原始消息、主数据、质量仪表盘", "覆盖率≥99.9%；可重复重建盘口；时钟误差可测", "$0"),
        ("1 回测与基线", "第 7–12 周", "事件驱动回测；2 个 P0 策略；实验注册", "样本外扣费后为正；无明显泄漏；PnL 可逐笔解释", "$0"),
        ("2 影子交易", "第 13–18 周", "实时报价但不下单；模拟订单与平台盘口对比", "连续 4–6 周；延迟/成交模型误差在阈值内；无 P0 故障", "$0"),
        ("3 接口金丝雀", "第 19–20 周", "最小金额真实订单、撤单、成交、结算、对账", "账户/账本完全一致；故障时能自动停机", "$100"),
        ("4 小资金验证", "第 21–26 周", "单策略、单平台、受限实盘", "≥300 已结算事件；最大回撤和执行偏差合格", "$300"),
        ("5 扩展", "6 个月后", "第二策略/平台；审计月报", "≥90 天稳定；≥1,000 结算记录；无资金/合规事件", "$1,000→$3,000"),
    ], [1200, 1100, 2800, 3160, 1100], font_size=7.8, center_cols=(1, 4))

    add_heading(doc, "4.3 3,000 美元资金部署方案", 2)
    add_para(doc, "3,000 美元应当是最高可损失金额，而不是第一笔入金。未部署资本不应留在交易平台。下面是条件式阶梯，不是收益承诺。")
    add_table(doc, ["层级", "入场条件", "账户资本", "单市场风险", "日止损", "升级条件"], [
        ("L0", "研究/模拟", "$0", "$0", "$0", "回测和影子门槛通过"),
        ("L1", "验证 API/账本", "$100", "$1–5", "$10", "20+ 完整生命周期无调账差异"),
        ("L2", "小资金金丝雀", "$300", "$5–15", "$15", "300+ 结算；至少 30 天"),
        ("L3", "初步实盘", "$1,000", "$10–30", "$30", "90 天；压力测试和对账通过"),
        ("L4", "研究资本上限", "$3,000", "$15–50", "$60", "需双方重新审批；不得自动升级"),
    ], [900, 2600, 1300, 1500, 1100, 1960], font_size=8.3, center_cols=(0, 2, 3, 4))
    add_para(doc, "对 0.99 附近的 Tail Sweep，不使用普通胜率作为仓位依据。应以“整笔本金可能损失”的压力场景计入单市场风险，并限制同时开放的结算窗口数量。")

    add_heading(doc, "4.4 六个月现金预算（不含人工）", 2)
    add_table(doc, ["项目", "精简方案", "建议方案", "备注"], [
        ("风险资本", "$300 起，最高 $3,000", "$3,000 预留但分阶段部署", "未部署部分不进入平台"),
        ("本地设备", "使用现有设备", "NVMe/UPS 补强 ¥3k–10k", "不追求 GPU；优先内存、NVMe、UPS"),
        ("云与存储", "$600–1,200 / 6个月", "$1,200–3,000 / 6个月", "单区域起步；按数据量扩容"),
        ("域名/监控/备份", "$100–300", "$300–600", "可优先采用开源组件"),
        ("法律与税务咨询", "不得省略", "¥5k–20k+", "应针对居住地、平台、资金和雇主关系"),
        ("总现金需求", "若只研究：约 $1k–2k＋咨询", "含 $3k 风险资本：约 $5k–8k＋咨询", "设备已有时的量级估算"),
    ], [2100, 2200, 2500, 2560], font_size=8.6)

    add_heading(doc, "4.5 开户、资金与实操风险", 2)
    for item in [
        "平台资格：平台允许注册不等于居住地法律允许交易；平台地理限制也会变化。Polymarket 明确禁止 VPN 绕过限制；Kalshi 对国际用户要求 KYC，并要求用户自行确认当地合法性。[S11][S13]",
        "实名与账户：禁止借用、代持、共用或远程控制另一方账户。交易账户、银行卡/卡片、钱包和税务主体必须一致。",
        "资金路径：链上 USDC、虚拟货币出入金、跨境汇款和银行卡入金分别涉及银行、平台、税务和外汇问题；逐项取得专业意见。",
        "密钥安全：硬件钱包或隔离签名、API 权限最小化、交易与提币权限分离、密钥轮换、应急撤销。",
        "平台风险：冻结、KYC 复核、费率变化、API 变更、限额、结算争议和平台故障都可能让策略与资金同时失效。",
        "记录义务：保存开户条款版本、每次入出金、成交、费用、返佣、链上交易和每日 NAV，以便调账和税务咨询。",
    ]:
        add_bullet(doc, item, bullet_num)

    add_heading(doc, "4.6 中国大陆与雇主合规的前置门槛", 2)
    add_callout(doc, "高风险提示", "人民银行等十部门银发〔2021〕237号目前仍被监管法规库标注为现行有效；其内容涉及虚拟货币、虚拟货币衍生品、境外平台向境内居民提供服务及相关资金活动。最高法也曾明确：庄家式、输者本金归网站的网络“二元期权”可被认定为赌博网站。[S21][S22] 本项目的 CLOB 对手盘结构并不等同于该案例，但不能因此排除适用其他金融、博彩、外汇或网络规定。", fill=LIGHT_RED, color=RED)
    add_para(doc, "因此本白皮书不判断某个平台或某种资金路径在中国大陆是否合法。阶段 −1 的产物必须包括：由熟悉虚拟资产、跨境交易与网络博彩边界的律师出具的针对性意见；若无法获得明确意见，则项目永久限制在公开数据研究、模拟和软件工程。")
    add_para(doc, "对象在量化公司任职还会触发个人账户交易、外部兼职、知识产权、竞业、保密和利益冲突规定。最稳妥做法是提交书面申报，只讨论公开的量化工程原则，并保留清洁室记录。禁止使用雇主电脑、代码、数据、交易账户、云资源或工作时间。")

    add_heading(doc, "5. 风险评估、时间评估与最终建议", 1)
    add_heading(doc, "5.1 风险矩阵", 2)
    add_table(doc, ["风险", "概率", "影响", "主要控制", "触发停机"], [
        ("法律/平台资格", "中–高", "极高", "律师意见、雇主申报、实名账户、禁绕过", "任何不确定性升级或平台警告"),
        ("结算定义错误", "中", "极高", "规则版本、双人复核、模糊市场禁交易", "规则源不一致/无法解析"),
        ("回测过拟合", "高", "高", "冻结样本外、全实验登记、PBO/压力测试", "收益集中或样本外反转"),
        ("延迟/断流", "高", "高", "多源心跳、时钟监控、fail closed", "stale feed、时钟偏差超限"),
        ("重复/错误下单", "中", "极高", "幂等 ID、状态机、单市场限额、总开关", "内部账本与平台不一致"),
        ("平台/链上最终性", "中", "高", "MATCHED≠CONFIRMED；重试与调账", "RETRYING/FAILED 累积"),
        ("尾部全损", "中", "高", "按本金损失计风险、并发窗口限制", "单日亏损/异常结算"),
        ("资金冻结/出金", "中", "高", "小余额、定期出金演练、资金分离", "KYC 复核或提现失败"),
        ("家庭协作", "中", "中", "职责分离、双人审批、固定复盘", "争议时默认不交易"),
        ("时间与倦怠", "高", "中", "限制范围、每周时长、明确 kill criteria", "连续 4 周无法交付里程碑"),
    ], [1700, 900, 900, 3900, 1960], font_size=7.9, center_cols=(1, 2))

    add_heading(doc, "5.2 时间评估", 2)
    add_table(doc, ["目标", "乐观", "基准", "保守"], [
        ("可靠数据底座", "4 周", "6–8 周", "12 周"),
        ("可审计事件驱动回测", "8 周", "12–16 周", "24 周"),
        ("连续影子交易", "第 3 月", "第 4–5 月", "第 6–8 月"),
        ("首笔受控实盘", "第 4 月", "第 5–6 月", "合规或模型不通过则不做"),
        ("判断策略是否可继续", "6 个月", "9–12 个月", "可能始终无法证明"),
    ], [3200, 1800, 2200, 2160], center_cols=(1, 2, 3))

    add_heading(doc, "5.3 Kill Criteria（停止标准）", 2)
    for item in [
        "无法取得明确的平台资格或居住地法律意见；任何方案依赖 VPN、代持或虚假 KYC。",
        "对象所在公司不允许该类个人交易/外部项目，或无法建立清洁室边界。",
        "连续 8–12 周仍不能实现可重复订单簿重建和逐笔 PnL 调账。",
        "样本外收益扣除费用与最坏合理执行假设后为负，或超过 30% 利润来自单一事件/单日。",
        "影子交易与回测的成交/滑点偏差长期不可解释。",
        "发生密钥泄露、未授权订单、账本不一致、无法撤单或资金出入异常。",
        "任何一方认为家庭财务、工作或关系成本超过研究价值。",
    ]:
        add_bullet(doc, item, bullet_num, color=RED)

    add_heading(doc, "5.4 条件式 GO / NO-GO", 2)
    add_table(doc, ["决策", "结论", "理由"], [
        ("是否值得研究", "GO", "赛道增长、数据开放、系统问题有技术含量，团队能力互补"),
        ("是否适合清洁室复刻", "GO", "基于公开 API 和独立代码可形成可验证能力；禁止复制私有资产"),
        ("是否立即开户实盘", "NO-GO", "法律、雇主、平台资格和资金路径尚未完成审查"),
        ("是否首日投入 $3,000", "NO-GO", "缺少独立回测、影子交易、对账和风险控制"),
        ("是否以 6 个月为首期", "GO", "足够验证工程和流程；不承诺产生可持续 Alpha"),
        ("是否主攻极低延迟", "暂缓", "资本小、运维重；先建立可离线验证的结构/lead-lag 基线"),
    ], [2300, 1300, 5760], center_cols=(1,))
    add_callout(doc, "最终建议", "把项目立项为“六个月、公开数据、清洁室实现、先模拟后实盘”的家庭研究计划。第一阶段成功标准不是赚钱，而是：系统可复现、账目可审计、风险可停止、合规可解释。若这些成立，再讨论 3,000 美元全部部署。", fill=LIGHT_BLUE, color=DARK_BLUE)

    add_heading(doc, "附录 A：与对象的首次对齐会议议程", 1)
    for item in [
        "确认项目性质：家庭自营研究，不是公司、不募资、不代客、不使用朋友或雇主的非公开资产。",
        "逐条阅读对象的劳动合同、个人交易政策、外部兼职/IP/竞业与申报要求。",
        "讨论量化公司中回测、OMS、风险、调账和上线审批的“验收原则”，不讨论雇主具体实现。",
        "选择唯一首期平台、唯一参考价格源、唯一市场类别；拒绝同时做多平台、多品类。",
        "决定 P0 Alpha：现货 lead-lag 与逻辑约束各一个；Tail Sweep 仅复盘。",
        "确定每周总投入时间、家庭预算、谁拥有一票停机权，以及何时复盘。",
        "列出必须向律师确认的事实：居住地、国籍/税务、平台、结算资产、入出金方式、自动化程度。",
    ]:
        add_number(doc, item, decimal_num)

    add_heading(doc, "附录 B：上线前检查清单", 1)
    checklist = [
        "法律意见与平台资格已存档；不依赖绕过地理限制。",
        "雇主书面申报/确认完成；清洁室记录建立。",
        "原始数据可回放，增量缺口、时钟和覆盖率有告警。",
        "市场规则版本化，模糊/争议市场默认不交易。",
        "回测包含费用、延迟、队列、部分成交、撤单和结算。",
        "所有实验留痕，测试集冻结，策略通过样本外和压力测试。",
        "OMS 幂等；MATCHED/MINED/CONFIRMED/FAILED 状态完整。",
        "单市场、策略、平台、账户和日内限额可由风险层强制执行。",
        "内部账本能与平台和链上逐笔调账；每日 NAV 锁定。",
        "密钥最小权限、备份、轮换和紧急撤销演练完成。",
        "影子交易连续 4–6 周；回测与实时偏差在阈值内。",
        "双方共同签字批准本次资本层级；未部署资本不在平台。",
    ]
    for item in checklist:
        add_bullet(doc, "□ " + item, bullet_num)

    add_heading(doc, "附录 C：资料来源", 1)
    add_para(doc, "资料访问日期：2026 年 8 月 17 日。平台规则、费率、地理限制和监管状态可能变化，实盘前必须重新核对。", size=9.5, italic=True, color=GRAY)
    sources = [
        ("S01", "SignalX Console（需授权访问）", "https://app.signalx.net/warehouse"),
        ("S02", "SignalX Research — Project Brief / Architecture & Strategies", "https://projectbrief.signalx.net/brief"),
        ("S03", "CFTC — Understanding Prediction Markets and Event Contracts", "https://www.cftc.gov/LearnandProtect/PredictionMarkets"),
        ("S04", "CFTC Staff Advisory 26-08 — Prediction Markets", "https://www.cftc.gov/csl/26-08/download"),
        ("S05", "CFTC 2026 Prediction Markets Rulemaking / Federal Register materials", "https://www.cftc.gov/media/14151/NPRM_PredictionMarkets060926/download"),
        ("S06", "Polymarket Docs — Market Data Overview", "https://docs.polymarket.com/market-data/overview"),
        ("S07", "Polymarket Docs — WebSocket User Channel / Order and Trade Lifecycle", "https://docs.polymarket.com/market-data/websocket/user-channel"),
        ("S08", "Polymarket Docs — Fees", "https://docs.polymarket.com/trading/fees"),
        ("S09", "Polymarket Docs — Maker Rebates", "https://docs.polymarket.com/market-makers/maker-rebates"),
        ("S10", "Polymarket Docs — Resolution", "https://docs.polymarket.com/concepts/resolution"),
        ("S11", "Polymarket Help — Geographic Restrictions", "https://help.polymarket.com/en/articles/13364163-geographic-restrictions"),
        ("S12", "Kalshi API — WebSocket Quick Start", "https://docs.kalshi.com/getting_started/quick_start_websockets"),
        ("S13", "Kalshi Help — International Access and Individual Signup", "https://help.kalshi.com/en/articles/14026044-can-i-trade-on-kalshi-from-outside-the-united-states"),
        ("S14", "Favorite–Longshot Bias research overview / primary literature entry", "https://pubs.aeaweb.org/doi/abs/10.1257/mic.2.1.58"),
        ("S15", "The Anatomy of a Decentralized Prediction Market: Microstructure Evidence from the Polymarket Order Book", "https://arxiv.org/abs/2604.24366"),
        ("S16", "Fill-Side Non-Retail Trading on Polymarket: Behavioral Tiers and Quote-Attribution Constraints", "https://arxiv.org/abs/2605.11640"),
        ("S17", "Bailey & López de Prado — How Backtest Overfitting Leads to False Discoveries", "https://escholarship.org/uc/item/9tq3327h"),
        ("S18", "Queuing Uncertainty of Limit Orders", "https://pubsonline.informs.org/doi/10.1287/mnsc.2023.03371"),
        ("S19", "Cloudflare R2 Pricing", "https://developers.cloudflare.com/r2/pricing/"),
        ("S20", "AWS EC2 Pricing", "https://aws.amazon.com/ec2/pricing/"),
        ("S21", "中国人民银行等十部门：银发〔2021〕237号", "https://www.pbc.gov.cn/tiaofasi/144941/3581332/4348658/index.html"),
        ("S22", "最高人民法院：网络“二元期权”赌博性质相关典型案例说明", "https://www.court.gov.cn/fabu/xiangqing/283901.html"),
        ("S23", "CFTC Prediction Markets Customer Fact Sheet", "https://www.cftc.gov/sites/default/files/2026/04/PredictionMarketFactSheet.pdf"),
        ("S24", "The Block — July 2026 Kalshi/Polymarket combined volume estimate", "https://www.theblock.co/amp/post/410382/kalshi-polymarket-volume-july"),
    ]
    add_table(doc, ["编号", "来源", "URL"], sources, [900, 3850, 4610], font_size=7.4, center_cols=(0,))

    add_heading(doc, "附录 D：关键假设与版本说明", 1)
    for item in [
        "项目为两名自然人的家庭内部研究，不向第三方募集资金、代客交易、销售信号或提供平台服务。",
        "所有实现基于公开资料、公开 API 和自主开发；不复制 SignalX 或任何雇主的代码、配置和私有数据。",
        "IT 成本为量级估算，不包含人工机会成本、税务、法律服务的最终报价和不可预见的平台费用。",
        "3,000 美元为最高风险资本设想，不代表建议投资金额，也不代表项目具备正收益。",
        "v0.1 重点是认知对齐；下一版本应在对象与法律专业人士反馈后更新平台范围、账户路径、P0 Alpha 与验收阈值。",
    ]:
        add_bullet(doc, item, bullet_num)

    add_para(doc, "— 文档结束 —", size=9.5, italic=True, color=GRAY,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=24, after=0, line=1.0)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
